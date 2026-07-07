"""
build_docx.py - Construtor de documentos Word .docx com identidade visual
da marca configurada em /opt/MAIA/brand/brand.json (via brand_loader).

Recebe um AST simples (lista de blocos parseados de markdown) e monta o .docx
aplicando a paleta, tipografia, capa, sumario, header e footer da marca.

Uso programatico:
    from build_docx import DocxBuilder
    b = DocxBuilder(title="Meu Doc", subtitle="...")  # autor vem do brand_loader
    b.add_heading(1, "Introducao")
    b.add_paragraph("Texto do paragrafo")
    b.save("/path/saida.docx")
"""
from __future__ import annotations

import os
import re
from datetime import datetime
from typing import List, Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import sys
sys.path.insert(0, "/opt/MAIA")
from brand_loader import footer_handle, brand_name, slogan_or_blank, website_or_blank, owner_name, get, colors


# ============================================================================
# IDENTIDADE VISUAL - paleta e tipografia (default da skill).
# As cores podem vir do sistema central via brand_loader.colors(); se vazio,
# cai no default abaixo (PALETA_DEFAULT_EXISTENTE).
# ============================================================================
COLOR_BLACK = RGBColor(0x0A, 0x0A, 0x0A)
COLOR_TEAL = RGBColor(0x3A, 0x9E, 0x9C)
COLOR_GOLD = RGBColor(0xC9, 0xA9, 0x6E)
COLOR_OFFWHITE = RGBColor(0xF5, 0xF3, 0xEE)
COLOR_GREY = RGBColor(0x6E, 0x6E, 0x6E)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_TEAL = "3A9E9C"
HEX_GOLD = "C9A96E"
HEX_OFFWHITE = "F5F3EE"
HEX_CODE_BG = "EFEEEA"

# Fontes
FONT_DISPLAY = "Playfair Display"
FONT_BODY = "Plus Jakarta Sans"
FONT_MONO = "JetBrains Mono"


# ============================================================================
# Helpers XML
# ============================================================================
def _make_element(tag, **kwargs):
    el = OxmlElement(tag)
    for k, v in kwargs.items():
        el.set(qn(k), str(v))
    return el


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = _make_element("w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": hex_color})
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex="C9A96E", size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = _make_element("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = _make_element(f"w:{edge}", **{"w:val": "single", "w:sz": size, "w:space": "0", "w:color": color_hex})
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_paragraph_border_left(paragraph, color_hex="C9A96E", size=24):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = _make_element("w:pBdr")
    left = _make_element("w:left", **{"w:val": "single", "w:sz": size, "w:space": "12", "w:color": color_hex})
    p_bdr.append(left)
    p_pr.append(p_bdr)


def set_paragraph_shading(paragraph, hex_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = _make_element("w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": hex_color})
    p_pr.append(shd)


def add_horizontal_line(paragraph, color_hex="C9A96E"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = _make_element("w:pBdr")
    bottom = _make_element("w:bottom", **{"w:val": "single", "w:sz": 8, "w:space": "1", "w:color": color_hex})
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_field(paragraph, instr_text):
    run = paragraph.add_run()
    fld_char_begin = _make_element("w:fldChar", **{"w:fldCharType": "begin"})
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_char_separate = _make_element("w:fldChar", **{"w:fldCharType": "separate"})
    placeholder = OxmlElement("w:t")
    placeholder.text = ""
    fld_char_end = _make_element("w:fldChar", **{"w:fldCharType": "end"})
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_separate)
    run._r.append(placeholder)
    run._r.append(fld_char_end)
    return run


def add_page_break_run(paragraph):
    run = paragraph.add_run()
    br = _make_element("w:br", **{"w:type": "page"})
    run._r.append(br)


# ============================================================================
# Builder
# ============================================================================
class DocxBuilder:
    def __init__(
        self,
        title,
        subtitle="",
        author=None,
        date_str=None,
        contact_email=None,
        contact_site=None,
        brand_label=None,
    ):
        self.title = title
        self.subtitle = subtitle
        self.author = author or owner_name() or "Autor"
        self.date_str = date_str or datetime.now().strftime("%d/%m/%Y")
        self.contact_email = contact_email if contact_email is not None else get("email", "")
        self.contact_site = contact_site if contact_site is not None else website_or_blank()
        self.brand_label = brand_label if brand_label is not None else (brand_name() or "")

        self.doc = Document()
        self._setup_page()
        self._setup_styles()
        self._build_header_footer()
        self._build_cover()
        self._build_toc()

    # ----- setup -----
    def _setup_page(self):
        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            section.different_first_page_header_footer = True
            section.header_distance = Cm(1.2)
            section.footer_distance = Cm(1.2)

    def _setup_styles(self):
        styles = self.doc.styles

        normal = styles["Normal"]
        normal.font.name = FONT_BODY
        normal.font.size = Pt(11)
        normal.font.color.rgb = COLOR_BLACK
        pf = normal.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.space_after = Pt(8)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        h1 = styles["Heading 1"]
        h1.font.name = FONT_DISPLAY
        h1.font.size = Pt(22)
        h1.font.color.rgb = COLOR_GOLD
        h1.font.bold = True
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(12)
        h1.paragraph_format.keep_with_next = True

        h2 = styles["Heading 2"]
        h2.font.name = FONT_BODY
        h2.font.size = Pt(16)
        h2.font.color.rgb = COLOR_TEAL
        h2.font.bold = True
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(8)
        h2.paragraph_format.keep_with_next = True

        h3 = styles["Heading 3"]
        h3.font.name = FONT_BODY
        h3.font.size = Pt(14)
        h3.font.color.rgb = COLOR_BLACK
        h3.font.bold = True
        h3.paragraph_format.space_before = Pt(14)
        h3.paragraph_format.space_after = Pt(6)
        h3.paragraph_format.keep_with_next = True

        h4 = styles["Heading 4"]
        h4.font.name = FONT_BODY
        h4.font.size = Pt(12)
        h4.font.color.rgb = COLOR_BLACK
        h4.font.bold = True
        h4.font.italic = True
        h4.paragraph_format.space_before = Pt(10)
        h4.paragraph_format.space_after = Pt(4)

        existing = [s.name for s in styles]

        if "CoverTitle" not in existing:
            cover_title = styles.add_style("CoverTitle", WD_STYLE_TYPE.PARAGRAPH)
            cover_title.font.name = FONT_DISPLAY
            cover_title.font.size = Pt(36)
            cover_title.font.color.rgb = COLOR_GOLD
            cover_title.font.italic = True
            cover_title.font.bold = False
            cover_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cover_title.paragraph_format.space_after = Pt(8)
            cover_title.paragraph_format.line_spacing = 1.1

        if "CoverSubtitle" not in existing:
            cover_sub = styles.add_style("CoverSubtitle", WD_STYLE_TYPE.PARAGRAPH)
            cover_sub.font.name = FONT_BODY
            cover_sub.font.size = Pt(16)
            cover_sub.font.color.rgb = COLOR_TEAL
            cover_sub.paragraph_format.space_after = Pt(20)

        if "CoverBrand" not in existing:
            cover_brand = styles.add_style("CoverBrand", WD_STYLE_TYPE.PARAGRAPH)
            cover_brand.font.name = FONT_BODY
            cover_brand.font.size = Pt(11)
            cover_brand.font.color.rgb = COLOR_GOLD
            cover_brand.font.bold = True

        if "CoverMeta" not in existing:
            cover_meta = styles.add_style("CoverMeta", WD_STYLE_TYPE.PARAGRAPH)
            cover_meta.font.name = FONT_BODY
            cover_meta.font.size = Pt(10)
            cover_meta.font.color.rgb = COLOR_GREY

        if "Callout" not in existing:
            callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
            callout.font.name = FONT_BODY
            callout.font.size = Pt(11)
            callout.font.italic = True
            callout.font.color.rgb = COLOR_BLACK
            callout.paragraph_format.left_indent = Cm(0.4)
            callout.paragraph_format.space_before = Pt(8)
            callout.paragraph_format.space_after = Pt(8)
            callout.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if "CodeBlock" not in existing:
            cb = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
            cb.font.name = FONT_MONO
            cb.font.size = Pt(10)
            cb.font.color.rgb = COLOR_BLACK
            cb.paragraph_format.space_before = Pt(2)
            cb.paragraph_format.space_after = Pt(2)
            cb.paragraph_format.line_spacing = 1.25
            cb.paragraph_format.left_indent = Cm(0.3)

    # ----- header / footer -----
    def _build_header_footer(self):
        section = self.doc.sections[0]

        first_header = section.first_page_header
        if first_header.paragraphs:
            first_header.paragraphs[0].text = ""
        first_footer = section.first_page_footer
        if first_footer.paragraphs:
            p = first_footer.paragraphs[0]
            p.text = ""
            run = p.add_run(self.date_str)
            run.font.name = FONT_BODY
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_GREY
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        header = section.header
        h_para = header.paragraphs[0]
        h_para.text = ""
        header_text = (self.brand_label + "  -  " + self.title) if self.brand_label else self.title
        run = h_para.add_run(header_text)
        run.font.name = FONT_BODY
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_TEAL
        run.font.bold = True
        h_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_horizontal_line(h_para, color_hex=HEX_TEAL)

        footer = section.footer
        f_para = footer.paragraphs[0]
        f_para.text = ""
        f_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        contact_parts = [x for x in (self.contact_email, self.contact_site) if x]
        contact_prefix = ("  -  ".join(contact_parts) + "  -  Pagina ") if contact_parts else "Pagina "
        run = f_para.add_run(contact_prefix)
        run.font.name = FONT_BODY
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_GREY
        run_page = add_field(f_para, "PAGE")
        run_page.font.name = FONT_BODY
        run_page.font.size = Pt(9)
        run_page.font.color.rgb = COLOR_GREY

    # ----- capa -----
    def _build_cover(self):
        p = self.doc.add_paragraph(style="CoverBrand")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Cm(3)
        run = p.add_run(self.brand_label)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_GOLD
        run.font.bold = True

        p2 = self.doc.add_paragraph(style="CoverMeta")
        run = p2.add_run("por " + (owner_name() or self.author))
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_GREY
        add_horizontal_line(p2, color_hex=HEX_GOLD)

        for _ in range(6):
            self.doc.add_paragraph()

        p_title = self.doc.add_paragraph(style="CoverTitle")
        run = p_title.add_run(self.title)
        run.font.name = FONT_DISPLAY
        run.font.italic = True
        run.font.color.rgb = COLOR_GOLD
        run.font.size = Pt(36)

        if self.subtitle:
            p_sub = self.doc.add_paragraph(style="CoverSubtitle")
            run = p_sub.add_run(self.subtitle)
            run.font.name = FONT_BODY
            run.font.size = Pt(16)
            run.font.color.rgb = COLOR_TEAL

        sep = self.doc.add_paragraph()
        add_horizontal_line(sep, color_hex=HEX_GOLD)

        for _ in range(8):
            self.doc.add_paragraph()
        p_auth = self.doc.add_paragraph()
        run = p_auth.add_run(self.author)
        run.font.name = FONT_BODY
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_BLACK

        p_date = self.doc.add_paragraph()
        run = p_date.add_run(self.date_str)
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_GREY

        end = self.doc.add_paragraph()
        add_page_break_run(end)

    # ----- sumario -----
    def _build_toc(self):
        p = self.doc.add_paragraph()
        run = p.add_run("Sumario")
        run.font.name = FONT_DISPLAY
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = COLOR_GOLD
        p.paragraph_format.space_after = Pt(16)

        p_toc = self.doc.add_paragraph()
        add_field(p_toc, 'TOC \\o "1-2" \\h \\z \\u')

        info = self.doc.add_paragraph()
        run = info.add_run(
            "(Se o sumario aparecer vazio ao abrir, clique com o direito sobre ele e escolha 'Atualizar campo'.)"
        )
        run.font.name = FONT_BODY
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_GREY
        run.font.italic = True

        end = self.doc.add_paragraph()
        add_page_break_run(end)

    # ----- API publica -----
    def add_heading(self, level, text):
        level = max(1, min(level, 4))
        p = self.doc.add_paragraph(style="Heading " + str(level))
        p.add_run(text)
        return p

    def add_paragraph(self, text, style=None):
        p = self.doc.add_paragraph(style=style or "Normal")
        self._add_inline_runs(p, text)
        return p

    def _add_inline_runs(self, paragraph, text):
        pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
        parts = pattern.split(text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**") and len(part) > 4:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                run.font.name = FONT_BODY
            elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
                run.font.name = FONT_BODY
            elif part.startswith("`") and part.endswith("`") and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.font.name = FONT_MONO
                run.font.size = Pt(10)
            else:
                run = paragraph.add_run(part)
                run.font.name = FONT_BODY
        return paragraph

    def add_bullet(self, text, level=0):
        p = self.doc.add_paragraph()
        run = p.add_run("•  ")
        run.font.name = FONT_BODY
        run.font.color.rgb = COLOR_GOLD
        run.font.bold = True
        self._add_inline_runs(p, text)
        p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_numbered(self, text, number):
        p = self.doc.add_paragraph()
        run = p.add_run(str(number) + ".  ")
        run.font.name = FONT_BODY
        run.font.color.rgb = COLOR_GOLD
        run.font.bold = True
        self._add_inline_runs(p, text)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_quote(self, text):
        p = self.doc.add_paragraph(style="Callout")
        set_paragraph_shading(p, HEX_OFFWHITE)
        add_paragraph_border_left(p, color_hex=HEX_GOLD, size=24)
        run = p.add_run(text)
        run.font.italic = True
        run.font.name = FONT_BODY
        run.font.size = Pt(11)
        return p

    def add_code_block(self, code, language=""):
        lines = code.splitlines() or [""]
        for i, line in enumerate(lines):
            p = self.doc.add_paragraph(style="CodeBlock")
            set_paragraph_shading(p, HEX_CODE_BG)
            run = p.add_run(line if line else " ")
            run.font.name = FONT_MONO
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_BLACK
            if i > 0:
                p.paragraph_format.space_before = Pt(0)
            if i < len(lines) - 1:
                p.paragraph_format.space_after = Pt(0)

    def add_table(self, headers, rows):
        if not headers:
            return
        cols = len(headers)
        table = self.doc.add_table(rows=1 + len(rows), cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            cell = hdr_cells[i]
            set_cell_shading(cell, HEX_TEAL)
            set_cell_borders(cell, color_hex=HEX_GOLD, size=4)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.name = FONT_BODY
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_WHITE
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

        for r_idx, row in enumerate(rows):
            tr_cells = table.rows[r_idx + 1].cells
            zebra = r_idx % 2 == 1
            for c_idx, val in enumerate(row):
                if c_idx >= cols:
                    break
                cell = tr_cells[c_idx]
                if zebra:
                    set_cell_shading(cell, HEX_OFFWHITE)
                set_cell_borders(cell, color_hex="C9C5BA", size=4)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.name = FONT_BODY
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_BLACK
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)

        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

    def add_horizontal_rule(self):
        p = self.doc.add_paragraph()
        add_horizontal_line(p, color_hex=HEX_GOLD)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    def add_page_break(self):
        p = self.doc.add_paragraph()
        add_page_break_run(p)

    def save(self, path):
        os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
        self.doc.save(path)
        return path
